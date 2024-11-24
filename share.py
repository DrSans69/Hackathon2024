from flask import Flask, jsonify, request
from flask_cors import CORS
import os
from werkzeug.utils import secure_filename
from app import llm

from docx import Document
from pdfminer.high_level import extract_text


import aspose.words as aw
import re


def pdf_to_html(temp_file):
    doc = aw.Document(temp_file)
    doc.save("output2.html")


    with open("output2.html", "r", encoding="utf-8") as file:
        html_string = file.read()
    html_string = re.sub(r'<meta name="generator" content="Aspose.Words for Python via .NET [^"]*" ?/?>', '', html_string)

    # Remove Aspose evaluation messages and links
    html_string = re.sub(
        r'<p[^>]*><span[^>]*>Created with an evaluation copy of Aspose\.Words\..*?</a></p>',
        '',
        html_string,
        flags=re.DOTALL
    )

    # Remove footer messages mentioning Aspose
    html_string = re.sub(
        r'<div[^>]*-aw-headerfooter-type:footer-primary[^>]*>.*?Created with Aspose\.Words[^<]*</span></p></div>',
        '',
        html_string,
        flags=re.DOTALL
    )

    # Clean up extra spaces or empty tags that may result from removal
    html_string = re.sub(r'\s*\n\s*', '\n', html_string)
    html_string = re.sub(r'<p[^>]*>\s*</p>', '', html_string)  # Remove empty <p> tags
    html_string = re.sub(r'<div[^>]*>\s*</div>', '', html_string)
    os.remove("output2.001.png")
    os.remove("output2.html")
    # Remove empty <div> tags

    # Output the cleaned HTML string
    return(html_string)  # For PDF to text conversion


def docx_to_html(temp_path):
    doc = aw.Document(temp_path)
    doc.save("Output.html")

    with open("output.html", "r", encoding="utf-8") as file:
        html_string = file.read()
    html_string = re.sub(r'<meta name="generator" content="Aspose.Words for Python via .NET [^"]*" ?/?>', '', html_string)

    # Remove Aspose evaluation messages and links
    html_string = re.sub(
        r'<p[^>]*><span[^>]*>Created with an evaluation copy of Aspose\.Words\..*?</a></p>',
        '',
        html_string,
        flags=re.DOTALL
    )

    # Remove footer messages mentioning Aspose
    html_string = re.sub(
        r'<div[^>]*-aw-headerfooter-type:footer-primary[^>]*>.*?Created with Aspose\.Words[^<]*</span></p></div>',
        '',
        html_string,
        flags=re.DOTALL
    )

    # Clean up extra spaces or empty tags that may result from removal
    html_string = re.sub(r'\s*\n\s*', '\n', html_string)
    html_string = re.sub(r'<p[^>]*>\s*</p>', '', html_string)  # Remove empty <p> tags
    html_string = re.sub(r'<div[^>]*>\s*</div>', '', html_string)  # Remove empty <div> tags

    # Output the cleaned HTML string
    os.remove("output.001.png")
    os.remove("Output.html")

    return(html_string)


app = Flask(__name__)
CORS(app)


# Supported file extensions for processing
ALLOWED_EXTENSIONS = {'docx', 'pdf', 'txt'}

# Helper function to check file extension
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Convert DOCX to text
def docx_to_text(file_path):
    doc = Document(file_path)
    return '\n'.join(paragraph.text for paragraph in doc.paragraphs)

# Convert PDF to text
def pdf_to_text(file_path):
    return extract_text(file_path)

# Endpoint to convert file to HTML
@app.route('/api/data', methods=['POST'])
def get_data():
    input_text = request.form.get('text', '')
    standard = request.form.get('standard', '')

    # Check uploaded file for text or standard input
    if not input_text and 'file1' in request.files:
        file = request.files['file1']
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            temp_path = os.path.join('temp', filename)  # Save file temporarily

            file.save(temp_path)  # Save uploaded file
            ext = filename.rsplit('.', 1)[1].lower()

            # Extract text based on file extension
            if ext == 'docx':
                input_text = docx_to_html(temp_path)
            elif ext == 'pdf':
                input_text = pdf_to_html(temp_path)
            elif ext == 'txt':
                with open(temp_path, 'r', encoding='utf-8') as f:
                    input_text = f.read()

            os.remove(temp_path)  # Clean up temporary file

    if not standard and 'file2' in request.files:
        file = request.files['file2']
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            temp_path = os.path.join('temp', filename)

            file.save(temp_path)
            ext = filename.rsplit('.', 1)[1].lower()

            # Extract standard based on file extension
            if ext == 'docx':
                standard = docx_to_text(temp_path)
            elif ext == 'pdf':
                standard = pdf_to_text(temp_path)
            elif ext == 'txt':
                with open(temp_path, 'r', encoding='utf-8') as f:
                    standard = f.read()

            os.remove(temp_path)

    # Process the extracted data
    processed_data = llm.inspect(input_text, standard)

    # Save processed input as HTML
    file_path = "temp/output.html"
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(f"<html><body><pre>{input_text}</pre></body></html>")

    return jsonify(processed_data)

# Serve the saved HTML file
@app.route('/html')
def get_html():
    with open('temp/output.html', encoding="utf-8") as file:
        return file.read(), 200, {'Content-Type': 'text/html'}

if __name__ == '__main__':
    app.run(debug=True)